from flask import Blueprint, request, jsonify, session
from werkzeug.utils import secure_filename
from src.models.user import db, User, Document
import os
import uuid
import PyPDF2
import docx
from io import BytesIO

documents_bp = Blueprint('documents', __name__)

UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

# Create upload folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path, file_type):
    """Extract text content from uploaded files"""
    try:
        if file_type == 'text/plain':
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        
        elif file_type == 'application/pdf':
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        
        elif file_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            return "Unsupported file type for text extraction"
    
    except Exception as e:
        return f"Error extracting text: {str(e)}"

@documents_bp.route('/upload', methods=['POST'])
def upload_document():
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed. Supported types: txt, pdf, doc, docx'}), 400
        
        # Generate unique filename
        original_filename = secure_filename(file.filename)
        file_extension = original_filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4()}.{file_extension}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        # Save file
        file.save(file_path)
        
        # Extract text content
        content_text = extract_text_from_file(file_path, file.content_type)
        
        # Save document info to database
        document = Document(
            user_id=user_id,
            filename=unique_filename,
            original_filename=original_filename,
            file_path=file_path,
            file_type=file.content_type,
            content_text=content_text
        )
        
        db.session.add(document)
        db.session.commit()
        
        return jsonify({
            'message': 'Document uploaded successfully',
            'document': document.to_dict()
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@documents_bp.route('/', methods=['GET'])
def get_user_documents():
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        documents = Document.query.filter_by(user_id=user_id).order_by(Document.uploaded_at.desc()).all()
        
        return jsonify({
            'documents': [doc.to_dict() for doc in documents]
        }), 200
        
    except Exception as e:
        return jsonify({'error': f'Failed to get documents: {str(e)}'}), 500

@documents_bp.route('/<int:document_id>', methods=['GET'])
def get_document(document_id):
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        document = Document.query.filter_by(id=document_id, user_id=user_id).first()
        
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        return jsonify({
            'document': document.to_dict(),
            'content_preview': document.content_text[:500] + '...' if len(document.content_text) > 500 else document.content_text
        }), 200
        
    except Exception as e:
        return jsonify({'error': f'Failed to get document: {str(e)}'}), 500

@documents_bp.route('/<int:document_id>', methods=['DELETE'])
def delete_document(document_id):
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        document = Document.query.filter_by(id=document_id, user_id=user_id).first()
        
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        # Delete file from filesystem
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
        
        # Delete from database
        db.session.delete(document)
        db.session.commit()
        
        return jsonify({'message': 'Document deleted successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete document: {str(e)}'}), 500

